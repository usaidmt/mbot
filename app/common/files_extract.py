from fastapi import UploadFile
import docx
import pdfplumber
from io import BytesIO
from docx.table import _Cell
from docx.text.paragraph import Paragraph
import json

def read_docx(file: UploadFile) -> str:
    file_bytes = BytesIO(file.file.read())
    document = docx.Document(file_bytes)
    content = ""

    # Iterate over document elements in order
    for element in document.element.body:
        if element.tag.endswith('}p'):  # Paragraph
            para = Paragraph(element, document)
            content += para.text.strip() + "\n"

        elif element.tag.endswith('}tbl'):  # Table
            table = docx.table.Table(element, document)
            headers = []
            table_json = []

            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    headers = cells
                else:
                    row_dict = {headers[j]: cells[j] if j < len(cells) else "" for j in range(len(headers))}
                    table_json.append(row_dict)

            content += json.dumps(table_json, ensure_ascii=False) + "\n"

    return content.strip()

def oldest_read_docx(file: UploadFile) -> str:
    file_bytes = BytesIO(file.file.read())
    document = docx.Document(file_bytes)
    content = ""

    # Helper to recursively read paragraphs and tables from cells
    def read_cell(cell: _Cell) -> str:
        cell_text = ""
        for item in cell._element:
            if isinstance(item, Paragraph):
                cell_text += item.text + "\n"
        return cell_text

    # Iterate over document elements in order
    for element in document.element.body:
        if element.tag.endswith('}p'):  # Paragraph
            para = Paragraph(element, document)
            content += para.text.strip() + "\n"
        elif element.tag.endswith('}tbl'):  # Table
            table = docx.table.Table(element, document)
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                content += row_text + "\n"

    return content.strip()


def read_docx_old(file: UploadFile):
    file_bytes = BytesIO(file.file.read())  # Convert the file to a seekable byte stream
    doc = docx.Document(file_bytes)  # Open the document with python-docx
    content = ""
    for para in doc.paragraphs:
        content += para.text  # Extract text from each paragraph
    return content

# Function to read .pdf file
def read_pdf(file: UploadFile):
    text = ""
    with pdfplumber.open(file.file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text


# Function to read .txt file
def read_txt(file: UploadFile):
    return file.file.read().decode("utf-8")


# Function to handle file reading based on file extension
def read_file(file: UploadFile):
    file_extension = file.filename.split(".")[-1].lower()

    if file_extension == "txt":
        return read_txt(file)
    elif file_extension == "docx":
        return read_docx(file)
    elif file_extension == "pdf":
        return read_pdf(file)
    else:
        raise ValueError("Unsupported file type")

